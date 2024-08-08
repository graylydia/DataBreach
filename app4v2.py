import os
import re
import streamlit as st
import matplotlib.pyplot as plt
from docx import Document
from io import BytesIO
from transformers import AutoModelForCausalLM, AutoTokenizer
import torch

# Set Hugging Face token
os.environ["HUGGINGFACE_TOKEN"] = "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"

# Load the LLM model and tokenizer
@st.cache_resource
def load_model():
    model_name = "meta-llama/Meta-Llama-3-8B"
    tokenizer = AutoTokenizer.from_pretrained(model_name, use_auth_token=True, trust_remote_code=True)
    model = AutoModelForCausalLM.from_pretrained(model_name, use_auth_token=True, trust_remote_code=True)
    return model, tokenizer

model, tokenizer = load_model()

# Function to read text files from the directory
def read_text_files(directory):
    data = {}
    for filename in os.listdir(directory):
        if filename.endswith(".txt"):
            with open(os.path.join(directory, filename), 'r') as file:
                data[filename] = file.read().splitlines()
    return data

# Improved function to analyze passwords
def analyze_passwords(data):
    password_criteria = re.compile(
        r'^(?=.*[A-Z])(?=.*[a-z])(?=.*\d)(?=.*[@$!%*?&])[A-Za-z\d@$!%*?&]{12,}$')
    valid_passwords = [pwd for pwd in data if password_criteria.match(pwd)]
    return len(valid_passwords), valid_passwords

# Function to analyze MAC addresses
def analyze_mac_addresses(data):
    mac_pattern = re.compile(r"([0-9A-Fa-f]{2}[:]){5}([0-9A-Fa-f]{2})")
    valid_mac_addresses = [mac for mac in data if mac_pattern.match(mac)]
    return len(valid_mac_addresses), valid_mac_addresses

# Function to analyze IP addresses
def analyze_ip_addresses(data):
    ip_pattern = re.compile(r"(\d{1,3}\.){3}\d{1,3}")
    valid_ip_addresses = [ip for ip in data if ip_pattern.match(ip)]
    return len(valid_ip_addresses), valid_ip_addresses

# Function to analyze email addresses
def analyze_email_addresses(data):
    email_pattern = re.compile(r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+")
    valid_email_addresses = [email for email in data if email_pattern.match(email)]
    return len(valid_email_addresses), valid_email_addresses

# Function to analyze bank account numbers (simple validation)
def analyze_bank_accounts(data):
    bank_pattern = re.compile(r"\d{8,12}")
    valid_bank_accounts = [account for account in data if bank_pattern.match(account)]
    return len(valid_bank_accounts), valid_bank_accounts

# Function to analyze social security numbers
def analyze_ssns(data):
    ssn_pattern = re.compile(r"\d{3}-\d{2}-\d{4}")
    valid_ssns = [ssn for ssn in data if ssn_pattern.match(ssn)]
    return len(valid_ssns), valid_ssns

# Function to analyze financial information
def analyze_financial_info(data):
    # Placeholder for complex financial info analysis
    return len(data), data

# Function to analyze health care data
def analyze_health_care_data(data):
    # Placeholder for health care data analysis
    return len(data), data

# Function to determine data type based on content
def determine_data_type(data):
    if any(re.match(r"([0-9A-Fa-f]{2}[:]){5}([0-9A-Fa-f]{2})", entry) for entry in data):
        return "mac"
    elif any(re.match(r'^(?=.*[A-Z])(?=.*[a-z])(?=.*\d)(?=.*[@$!%*?&])[A-Za-z\d@$!%*?&]{12,}$', entry) for entry in data):
        return "password"
    elif any(re.match(r"(\d{1,3}\.){3}\d{1,3}", entry) for entry in data):
        return "ip"
    elif any(re.match(r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+", entry) for entry in data):
        return "email"
    elif any(re.match(r"\d{8,12}", entry) for entry in data):
        return "bank"
    elif any(re.match(r"\d{3}-\d{2}-\d{4}", entry) for entry in data):
        return "ssn"
    else:
        return "unknown"

# Function to analyze data with LLM
def analyze_with_llm(data):
    prompt = "Determine if the following data is related to healthcare or financial information:\n\n"
    for entry in data:
        prompt += entry + "\n"
    prompt += "\nAnswer:"

    inputs = tokenizer(prompt, return_tensors="pt", truncation=True, max_length=1024)
    outputs = model.generate(inputs.input_ids, max_new_tokens=100)
    response = tokenizer.decode(outputs[0], skip_special_tokens=True)

    if "healthcare" in response.lower():
        return "healthcare"
    elif "financial" in response.lower():
        return "financial"
    else:
        return "unknown"

# Function to generate a Word document report
def generate_report(analysis_results):
    doc = Document()
    doc.add_heading('Data Analysis Report', 0)

    for data_type, result in analysis_results.items():
        doc.add_heading(f'Type: {data_type}', level=1)
        doc.add_paragraph(f"Total Valid Entries: {result['total_count']}")
        doc.add_heading('Valid Data Samples:', level=2)
        for data in result['valid_data'][:5]:  # Show only first 5 samples
            doc.add_paragraph(data)
        doc.add_paragraph('Files Data was Found In:')
        for filename, count in result['file_counts'].items():
            doc.add_paragraph(f"{filename} [{count} valid entries]")

    return doc

# Streamlit app
st.title('Data Corpus Analysis')

if st.button('Analyze'):
    directory = "data_corpus"
    data_files = read_text_files(directory)

    analysis_results = {
        "mac": {"total_count": 0, "valid_data": [], "file_counts": {}},
        "password": {"total_count": 0, "valid_data": [], "file_counts": {}},
        "ip": {"total_count": 0, "valid_data": [], "file_counts": {}},
        "email": {"total_count": 0, "valid_data": [], "file_counts": {}},
        "bank": {"total_count": 0, "valid_data": [], "file_counts": {}},
        "ssn": {"total_count": 0, "valid_data": [], "file_counts": {}},
        "financial": {"total_count": 0, "valid_data": [], "file_counts": {}},
        "healthcare": {"total_count": 0, "valid_data": [], "file_counts": {}},
        "unknown": {"total_count": 0, "valid_data": [], "file_counts": {}}
    }

    for filename, data in data_files.items():
        data_type = determine_data_type(data)
        
        if data_type == "mac":
            count, valid_data = analyze_mac_addresses(data)
        elif data_type == "password":
            count, valid_data = analyze_passwords(data)
        elif data_type == "ip":
            count, valid_data = analyze_ip_addresses(data)
        elif data_type == "email":
            count, valid_data = analyze_email_addresses(data)
        elif data_type == "bank":
            count, valid_data = analyze_bank_accounts(data)
        elif data_type == "ssn":
            count, valid_data = analyze_ssns(data)
        else:
            # Use LLM to determine if the data is healthcare or financial
            data_type = analyze_with_llm(data)
            if data_type == "financial":
                count, valid_data = analyze_financial_info(data)
            elif data_type == "healthcare":
                count, valid_data = analyze_health_care_data(data)
            else:
                count, valid_data = 0, []

        analysis_results[data_type]["total_count"] += count
        analysis_results[data_type]["valid_data"].extend(valid_data)
        if count > 0:
            analysis_results[data_type]["file_counts"][filename] = count

    # Extracting counts for plotting
    data_types = list(analysis_results.keys())
    counts = [result["total_count"] for result in analysis_results.values()]

    # Plotting the bar chart
    fig, ax = plt.subplots()
    ax.barh(data_types, counts, color='skyblue')
    ax.set_xlabel('Number of Valid Entries')
    ax.set_title('Data Corpus Analysis Results')
    st.pyplot(fig)

    # Plotting the pie chart
    fig2, ax2 = plt.subplots()
    ax2.pie(counts, labels=data_types, autopct='%1.1f%%', startangle=90, colors=plt.cm.Paired(range(len(counts))))
    ax2.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle.
    ax2.set_title('Distribution of Data Types')
    st.pyplot(fig2)

    # Generate and display the Word document report
    doc = generate_report(analysis_results)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    st.download_button(
        label="Download Report",
        data=buffer,
        file_name="Data_Analysis_Report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
